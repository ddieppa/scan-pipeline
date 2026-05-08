from __future__ import annotations

from pathlib import Path

from app.extractors.common import ExtractionResult


def extract_docx(path: Path, max_paragraphs: int, max_table_rows: int) -> ExtractionResult:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX extraction") from exc

    doc = Document(str(path))
    chunks: list[str] = []
    for paragraph in doc.paragraphs[:max_paragraphs]:
        if paragraph.text.strip():
            chunks.append(paragraph.text.strip())

    for table in doc.tables:
        for row_index, row in enumerate(table.rows):
            if row_index >= max_table_rows:
                break
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                chunks.append(row_text)

    return ExtractionResult(
        text="\n".join(chunks),
        text_source="docx",
        pages_inspected=1,
        needs_ocr=False,
        metadata={"paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables)},
    )
